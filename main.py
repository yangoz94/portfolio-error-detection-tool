import pandas as pd
import streamlit as st
import requests
import nltk
import nltk.tokenize.punkt
from nltk.probability import FreqDist
from nltk.corpus import stopwords
from string import punctuation, digits
import language_tool_python
import functools as ft
import docx2txt
import pdfplumber
import ssl
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# The following lines are required for ensuring there is no error because of nltk downloader.
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context
nltk.download('punkt')
nltk.download('stopwords')

# PAGE CONFIGURATION
st.set_page_config(
    page_title="ESL Portfolio Tool",
    page_icon="🧊",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'About': "Oğuzhan Yangöz\n\n\nGWU\n\n\noguzhan.yangoz@fulbrightmail.org"
    }
)

# VARIABLES
items_string = ["-ORGANISATION-", "-TASKCONTENT-", "-GRAMMAR-", "-VOCABULARY-"]
punctuation = punctuation.replace("-", "") + '““' + "’" + digits
is_bad_rule = ["Use a third-person plural verb with ‘they’.",
               "Use a comma before ‘and’ if it connects two independent clauses (unless they are closely connected and short)."]


# FUNCTIONS
def get_total():
    """
    This function calculates the totalscore of the student by taking a sum of the subscores.
    """
    items = [st.session_state[item] for item in items_string]
    total = 0
    for item in items:
        total += int(item)
    del st.session_state['-TOTALSCORE-']
    st.session_state['-TOTALSCORE-'] = total
    return total


def reset_scores():
    """
    This function resets the all subscores and the totalscore of the student, setting them to zero.
    """
    for item in items_string:
        st.session_state[item] = 0
    st.session_state["-TOTALSCORE-"] = 0


def clear_input():
    """
    This function clears the input textbox.
    It is passed as a button parameter (on_click).
    """
    st.session_state["-TEXT-"] = ""


def check_input():
    """
    This function checks whether the input is empty or not.
    If empty, it prints an error message.
    """
    if st.session_state["-TEXT-"] == "":
        st.error("You did NOT enter a text!")


def get_word_count():
    """
    This function calculates the word count of the input text.
    """
    words = st.session_state["-TEXT-"].split()
    word_count = len(words)
    return word_count


def get_sentence_count(text):
    """
    This function calculates the sentence count of the input text.
    """
    sentences = nltk.tokenize.sent_tokenize(text)
    sentence_count = len(sentences)
    return sentence_count, sentences


def remove_punctuation(data):
    """
    This function removes punctuation from a string.
    """
    removed = data.translate(str.maketrans('', '', punctuation))
    return removed


def get_top10_words():
    """
    This function yields top 10 commonly used words, along with their frequencies, in the input text.
    """
    fdist = FreqDist()
    clean_text = remove_punctuation(st.session_state["-TEXT-"])
    words = clean_text.split()
    stopWords = stopwords.words("english")
    stopWords.extend(punctuation)
    for word in words:
        fdist[word.lower()] += 1
    for item in fdist:
        if item in stopWords:
            fdist.pop(item)  # pseudo variable
    wordlist, freqlist = [], []
    for word, freq in fdist.most_common(10):
        wordlist.append(word)
        freqlist.append(freq)
    data = {"Word": wordlist, "Frequency": freqlist}
    chart_data = pd.DataFrame(data)

    return st.subheader("Word Frequency Table Excluding Stopwords"), \
           st.dataframe(chart_data), \
           st.markdown("-----------------------------------------------------------------")


@ft.lru_cache(maxsize=None)
def cache_tool():
    """
    This function caches the language tool for faster processing after the initial run.
    """
    tool = language_tool_python.LanguageTool('en-US')
    return tool


def check_text():
    """
    This function checks the input text for language error matches.
    """
    matches = cache_tool().check(st.session_state["-TEXT-"])
    return matches


def divide_text_for_api_calls():
    """
    Unfortunately, the free API calls are restricted to 1500 character-long texts per call.
    This function divides the text into 1500-character long chunks to make multiple calls and join them.
    """
    count, chunk, result = 0, "", []
    for sentence in get_sentence_count(st.session_state["-TEXT-"])[1]:
        sentence = " " + sentence  # to eliminate redundant whitespace errors
        if count + len(sentence) < 1500:
            count += len(sentence)
            chunk += sentence
        else:
            result.append(chunk)
            count = 0
            chunk = ""
    result.append(chunk)
    return result


def remove_bad_rules(dict):
    """
    This function filters out the undesired or ineffective error messages in the output.
    """
    for key, value in list(dict.items()):
        if value in is_bad_rule:
            del dict[key]


def get_local_matches():
    """
    This function returns a dictionary with the matched text and the error message for local processing.
    """
    d = {}
    for match in check_text():
        d[match.matchedText] = match.message
    remove_bad_rules(d)
    return d


def replacement(myDict):
    """
    This function inserts error messages at the end of the errored words/sentences.
    """
    text = st.session_state["-TEXT-"]
    for error, message in myDict.items():
        text = text.replace(error, str(error) + " <" + str(message) + ">")
    return text


def get_api_call_matches():
    """
    This function makes API requests and returns a dictionary with the matched text and error messages.
    Internet connection required.
    """
    matches, api = [], {}
    for piece in divide_text_for_api_calls():
        parameters = {"text": piece, "language": "en-US", "level": "picky", "disabledRules": "WHITESPACE_RULE"}
        response = requests.post("https://api.languagetool.org/v2/check", params=parameters)
        matches += response.json()["matches"]
        print(response)
    for match in matches:
        api[match["sentence"]] = match["message"]
    remove_bad_rules(api)
    return api


def create_output_layout():
    """
    This function displays basic text and input elements.
    """
    st.subheader("Editable Output")
    st.text_area(
        label="Your editable output is shown below. Please feel free to make manual corrections before visualization if necessary .",
        key="-TEXTOUTPUT-",
        height=500)


def display_editable_output(function):
    """
    This function displays output.
    """
    st.session_state["-TEXTOUTPUT-"] = replacement(function)
    create_output_layout()
    st.button(label="Download as Word", key="-DOWNLOAD-", on_click=create_docx_file)


def display_stats():
    """
    This function displays input text statistics.
    """
    st.markdown("-----------------------------------------------------------------")
    st.subheader("Text Statistics")
    m1, m2, m3 = st.columns([1, .8, .5])
    m1.metric(label="Sentence Count", value=str(get_sentence_count(st.session_state["-TEXT-"])[0]) + "  sentence(s)", )
    m2.metric(label="Word Count", value=str(get_word_count()) + "  word(s)")
    m3.metric(label="Character Count", value=str(len(st.session_state["-TEXT-"])) + "  characters")
    st.markdown("-----------------------------------------------------------------")
    get_top10_words()


def get_locally_corrected_text():
    """
    This function returns automatically-corrected version of the input text.
    """
    display_stats()
    correctedText = cache_tool().correct(st.session_state["-TEXT-"])
    st.session_state["-TEXTOUTPUT-"] = correctedText
    create_output_layout()


def read_file():
    """
    This function updates the text input with the content of the uploaded file.
    """
    if uploaded_file is not None:
        if uploaded_file.type == "application/pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                pages = pdf.pages[0]
                st.session_state["-TEXT-"] = pages.extract_text()
        else:
            raw_text = docx2txt.process(uploaded_file)
            del st.session_state["-TEXT-"]
            st.session_state["-TEXT-"] = raw_text
    else:
        st.error("Oops, are you sure you uploaded a file? ")


def create_docx_file():
    """
    This function creates a final report with all input and output details of the relevant student.
    """
    filename = st.session_state["-STUDENT-"] + " " + st.session_state["-ESSAY-"] + " " + \
               st.session_state["-DRAFT-"] + ".docx"

    records = (
        ("Task Content", str(st.session_state["-TASKCONTENT-"])),
        ("Grammar", str(st.session_state["-GRAMMAR-"])),
        ("Vocabulary", str(st.session_state["-VOCABULARY-"])),
        ("Organisation", str(st.session_state["-ORGANISATION-"])),
        ("Total Score", str(st.session_state["-TOTALSCORE-"]))
    )
    document = Document()
    section = document.sections[0]
    header = section.header
    headert = header.paragraphs[0]
    headert.text = "Instructor: " + st.session_state["-TEACHER-"]

    document.add_heading(
        st.session_state["-CLASSS-"] + " - " + st.session_state["-STUDENT-"] + " - " + st.session_state[
            "-ESSAY-"] + " - " + st.session_state["-DRAFT-"], 0)

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Grading Categories'
    hdr_cells[1].text = 'Scores'

    for category, score in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(category)
        row_cells[1].text = str(score)

    document.add_heading("Original / Submitted Text", 2)
    p1 = document.add_paragraph(st.session_state["-TEXT-"])
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    document.add_heading("Graded Text /  Feedback", 2)
    p2 = document.add_paragraph(st.session_state["-TEXTOUTPUT-"])
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    document.add_page_break()
    document.save(filename)
    st.success('The file has successfully been saved to your computer!\n\n '
               'Please search it on your computer as: ' + filename)


# MAIN WINDOW DESIGN
st.title('Portfolio Tool for English Language Teaching')
st.markdown("This tool aims to help English Language Instructors with their portfolio grading."
            "\nPlease note that the tool uses open source libraries only, and error detector and correction features might NOT be 100% accurate."
            "\n\nPlease copy and paste the student essay or upload it as a word or pdf file for text processing.")
st.text_area(
    label="After copying the essay, click one of the three buttons; Check Online, Check Locally or Correct Locally.",
    key="-TEXT-", height=500)
uploaded_file = st.file_uploader("Upload a file (Optional)", key="-FILE-",
                                 help="Please choose a file from your computer to upload a text.",
                                 type=["doc", "docx", "pdf"])

col1, col2, col3, col4 = st.columns([0.7, 3.5, 3, 1])

with col1:
    readB = st.button("Read File", on_click=read_file)
    clearB = st.button("Clear All", on_click=clear_input)
    ballonB = st.button("Balloons", on_click=st.balloons)

with col2:
    st.caption(
            'If you uploaded a file, click this button to copy and paste the content of your file to the box above.')
    st.caption('It erases all the input and output. <br>File itself and the sidebar will NOT be affected.', unsafe_allow_html=True)
    st.caption('Do you love balloons? If yes, go ahead!')

with col3:
    st.caption(
        'It checks the text for errors online in sentence-level. It is faster than local processing. Internet connection required.')
    st.caption('It checks the text for errors locally in word-level and inserts error messages; it might be slower than "Check Online.')
    st.caption('It corrects the text online through an API and outputs the corrected version of the input. It might take a while.')
with col4:
    apiB = st.button("Check Online", on_click=check_input)
    localB = st.button("Check Locally", on_click=check_input)
    localCorrectB = st.button("Correct Locally", on_click=check_input)

# SIDEBAR DESIGN
st.sidebar.header("Data Entry Section")
st.sidebar.markdown("Please enter the requested data before you copy a student essay and process it.")
st.sidebar.text_input("Teacher Name", help="Please enter your name e.g. Oğuzhan Yangöz", key="-TEACHER-")
st.sidebar.text_input("Class Name", help="Please enter the class name e.g. L-101.", key="-CLASSS-")
st.sidebar.text_input("Student Name", help="Please enter the student's name e.g. Oğuzhan.", key="-STUDENT-")
st.sidebar.selectbox("Essay Type", help="Please choose an essay type from the menu.", key="-ESSAY-",
                     options=(
                         "Opinion Essay", "Argumentative Essay", "Compare and Contrast Essay",
                         "Cause and Effect Essay", "Other"))
st.sidebar.selectbox("Draft No", help="Please choose a draft from the menu.", key="-DRAFT-",
                     options=("Draft 1", "Draft 2", "Draft 3"))

st.sidebar.header("Scoring Section")
st.sidebar.markdown(
    "Please enter a score between 0 and 25 for each subscore. Note that Total Score will automatically be calculated and displayed.")

subscore_settings = {"min_value": 0, "max_value": 25, "help": "Please enter a score between 0 and 25."}
st.sidebar.number_input("Use of English (25 pts)", key="-GRAMMAR-", **subscore_settings)
st.sidebar.number_input("Vocabulary (25 pts)", key="-VOCABULARY-", **subscore_settings)
st.sidebar.number_input("Task Content (25 pts)", key="-TASKCONTENT-", **subscore_settings)
st.sidebar.number_input("Organization (25 pts)", key="-ORGANISATION-", **subscore_settings)

calculateB = st.sidebar.button("Calculate Total", on_click=get_total, key="-CALCULATE-")
resetB = st.sidebar.button("Clear All Scores", on_click=reset_scores, key="-RESETSCORES-")
st.sidebar.number_input("Total (100 pts)", key="-TOTALSCORE-", min_value=0, max_value=100, )

# EVENTS
if resetB:
    st.sidebar.success("All scores have succesfully been reset!")
elif calculateB:
    st.sidebar.success("Total Score has succesfully been calculated!", )
elif localB or apiB or localCorrectB:
    if not st.session_state["-TEXT-"] == "":
        display_stats()
        if localB:
            display_editable_output(get_local_matches())
        elif apiB:
            display_editable_output(get_api_call_matches())
        elif localCorrectB:
            get_locally_corrected_text()


#column stylings that are not supported as built-in by streamlit=
st.markdown("""
    <style>
        [data-testid=column]:nth-of-type(1) [data-testid=stVerticalBlock]   {
            gap: 1.75rem;
        }
    </style>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
        [data-testid=column]:nth-of-type(2) [data-testid=stVerticalBlock]   {
            gap: 2.25rem;
        }
    </style>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
        [data-testid=column]:nth-of-type(4) [data-testid=stVerticalBlock]   {
            gap: 1.75rem;
        }
    </style>
    """, unsafe_allow_html=True)